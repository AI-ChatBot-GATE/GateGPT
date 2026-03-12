from docx import Document
import re

def edit_docx(input_path, output_path):
    doc = Document(input_path)
    
    # 1. Edit Paragraphs (Direct text replacement)
    replacements = {
        r"Group No\s*:\s*\d+": "Group No\t\t\t: 20",
        r"Title\s*:\s*Canteen Management System": "Title\t\t\t\t : GateGPT: AI-Powered GATE Preparation Platform",
        r"Product Name\s*:\s*Campus Bites": "Product Name\t\t : GateGPT",
        r"Full Stack Product referred for study :Ex : MIT CAMPUS DINING": "Full Stack Product referred for study : GateGPT Open Source",
        r"LINK: https://studentlife.mit.edu/dining/": "LINK: https://github.com/AI-ChatBot-GATE/GateGPT",
        r"Roll No \s*:\s*CB.SC.U4CSE23111": "Roll No \t\t\t:CB.SC.U4CSE23144",
        r"Name \s*:\s*B. GURU SAI PRASAD REDDY": "Name \t\t\t:Sandheep G S",
        r"Collaborator Name\s*:\s*.*": "Collaborator Name\t:"
    }
    
    for para in doc.paragraphs:
        for pattern, replacement in replacements.items():
            if re.search(pattern, para.text):
                # Preserving some formatting by only replacing content within runs if possible, 
                # but for simplicity in a template we'll replace the text.
                para.text = re.sub(pattern, replacement, para.text)

    # 2. Edit Table 0 (Metrics)
    if len(doc.tables) > 0:
        t0 = doc.tables[0]
        # Skip header Row 0
        # Row 1: Modules
        t0.cell(1, 2).text = "6"
        t0.cell(1, 3).text = "Auth, Chat, Learning, Notes, Practice, Analytics"
        # Row 2: Submodules
        t0.cell(2, 2).text = "18"
        t0.cell(2, 3).text = "Login, Signup, JWT Auth, Markdown Chat, LaTeX Solver, Subject Explorer, Revision Sheet Gen, Note CRUD, Bookmark Toggle, Score Tracking, Analytics Engine, Chat History, Persistent Sessions, Formula Rendering, Navigation Guard, Mobile Layout, API Proxy, Search"
        # Row 3: Master Forms
        t0.cell(3, 2).text = "4"
        t0.cell(3, 3).text = "Signup, Login, Create Note, Profile Edit"
        # Row 4: Transaction Forms
        t0.cell(4, 2).text = "3"
        t0.cell(4, 3).text = "Chat Message Box, Start Practice, Submit Response"
        # Row 5: Reports
        t0.cell(5, 2).text = "3"
        t0.cell(5, 3).text = "Practice History Table, Saved Notes Summary, Performance Report"
        # Row 6: Visualisations
        t0.cell(6, 2).text = "2"
        t0.cell(6, 3).text = "Subject Mastery Radar Chart, Weekly Score Trend"
        # Row 7: Master Tables
        t0.cell(7, 2).text = "3"
        t0.cell(7, 3).text = "Users, Subjects, Topics"
        # Row 8: Transaction Tables
        t0.cell(8, 2).text = "4"
        t0.cell(8, 3).text = "Chats, Messages, Notes, PracticeResults"
        # Row 9: Novelty
        t0.cell(9, 2).text = "AI-powered GATE tutoring with real-time LaTeX rendering, persistent cross-context study notes, automatic high-yield revision synthesis, and personalized performance analytics."
        t0.cell(9, 3).text = "Designed and implemented a persistent 'Chat-to-Note' integration, a high-yield AI revision generator using specialized LLM orchestration, and a LaTeX-native rendering engine for complex technical formulas."

    # 3. Edit Table 1 (Contributions)
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        # Row 1: Next.js (Formerly Node)
        t1.cell(1, 0).text = "Next.js"
        t1.cell(1, 1).text = "Built the responsive Next.js dashboard with Framer Motion, TailwindCSS, and integrated KaTeX for formula rendering."
        t1.cell(1, 4).text = "Developed a premium, glassmorphic UI that handles real-time AI streaming responses and technical formatting for engineering students."
        # Row 2: Express JS
        t1.cell(2, 1).text = "Developed the robust REST API using Express.js with JWT security and Gemini AI model orchestration."
        t1.cell(2, 4).text = "Designed modular routes for Auth, Chat, Note management, and AI tools, ensuring high availability and secure data transitions."
        # Row 3: MongoDB
        t1.cell(3, 1).text = "Designed schemas for Users, Chats, Notes, and Practice Analytics in MongoDB."
        t1.cell(3, 4).text = "Implemented optimized indexing for fast retrieval of sub-contextual notes and efficient tracking of student performance history."
        # Row 4: Integration
        t1.cell(4, 1).text = "Connected Next.js frontend with Express API using Axios with secure interceptors."
        t1.cell(4, 4).text = "Established a seamless data flow for the 'Save to Notes' feature, allowing chat insights to be instantly visible on subject syllabus pages."
        # Row 5: Innovation
        t1.cell(5, 1).text = "Implemented 'Chat to Note' workflow and Automated AI Revision Sheet generation."
        t1.cell(5, 4).text = "Innovated a dual-context learning model where AI tutor feedback is perpetually available as structured study material for exams."
        # Row 6: Testing
        t1.cell(6, 1).text = "Conducted end-to-end testing of the Auth flow, AI solving, and Note CRUD operations."
        t1.cell(6, 4).text = "Verified LaTeX rendering accuracy and API performance under simulated high-load exam preparation scenarios."
        # Row 7: Report
        t1.cell(7, 1).text = "Documented system architecture, database design, and AI model prompting strategies."
        t1.cell(7, 4).text = "Created comprehensive documentation for the backend Q&A and database relationships to facilitate future scaling."

    doc.save(output_path)
    print(f"File saved successfully to {output_path}")

if __name__ == "__main__":
    edit_docx("CB.SC.U4CSE23111Review 3_Template.docx", "CB.SC.U4CSE23144Review 3_GateGPT.docx")
