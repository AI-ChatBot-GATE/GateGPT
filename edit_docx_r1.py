from docx import Document
import re

def edit_docx_r1(input_path, output_path):
    doc = Document(input_path)
    
    # 1. Edit Paragraphs (Direct text replacement)
    replacements = {
        r"Project Title – .*": "Project Title – GateGPT: AI-Powered GATE Preparation Platform",
        r"Product Name .*": "Product Name GateGPT",
        r"Git Hub:\s*": "Git Hub: https://github.com/AI-ChatBot-GATE/GateGPT",
        r"https://github.com/Gurusaiprasadreddy/.*": "",
        r"Collaborator Details:.*": "Collaborator Details: ",
        r"Roll No \s*:\s*CB.SC.U4CSE23111": "Roll No \t\t\t:CB.SC.U4CSE23144",
        r"Name \s*:\s*B. GURU SAI PRASAD REDDY": "Name \t\t\t:Sandheep G S",
        r"Group No\s*:\s*17": "Group No\t\t\t: 20", # In case it's in a para
    }
    
    # Specific replacement for Abstract
    abstract_text = (
        "GateGPT is an intelligent full-stack preparation platform for the GATE exam, "
        "featuring an AI tutor (Gemini-based), real-time LaTeX formula rendering, and a "
        "persistent study-notes ecosystem. It bridges the gap between interactive AI assistance "
        "and formal study by allowing students to save AI insights directly into contextual "
        "subject-specific notes. With modules for revision generation, practice quizzes, and "
        "deep performance analytics, it provides a comprehensive digital environment for engineering students."
    )
    
    keywords_text = "Keywords: GATE, AI Tutor, Next.js, Express, MongoDB, LaTeX, Gemini AI, Persistent Notes, Study Analytics"

    for para in doc.paragraphs:
        # Simple replacements
        for pattern, replacement in replacements.items():
            if re.search(pattern, para.text):
                para.text = re.sub(pattern, replacement, para.text)
        
        # Abstract replacement logic (targeting P13-P16 based on our analysis)
        if "Campus Bites is a full-stack" in para.text:
            para.text = abstract_text
        elif "The platform serves three roles" in para.text:
            para.text = "The platform serves Students through an AI-powered tutoring interface, allowing them to manage study history, generate revision notes, and track subject mastery."
        elif "Its core innovation lies in a multi-layered" in para.text:
            para.text = "The core innovation lies in the AI-orchestration engine that handles technical solving with LaTeX, automated summarization for revision, and a cross-context note-saving system."
        elif "With 11 modules, 53 sub-modules" in para.text:
            para.text = "With 6 core modules including Auth, Chat, Learning, Notes, Practice, and Analytics, GateGPT delivers a specialized tech-stack for exam preparation."
        elif "Keywords:" in para.text:
            para.text = keywords_text
            
        # Module Explanations (Targeting Guru, Surya, Radha sections)
        if "Guru:" in para.text or "Surya:" in para.text or "Radha:" in para.text:
            para.text = "" # Clear names
        
        if "Module 1 — Management" in para.text:
            para.text = "Module 1 — Auth & Profile: Handles secure onboarding with JWT, role-based session management, and personalized user profile configuration."
        elif "Module 7 — User Management" in para.text:
            para.text = "Module 2 — AI Chat & Solver: The central tutoring hub where Gemini AI processes engineering queries and renders step-by-step solutions using LaTeX formatting."
        elif "Module 3 — Crew" in para.text:
            para.text = "Module 3 — Learning Explorer: Provides structured navigation through the GATE syllabus, allowing users to view subject topics and access related study materials."
        elif "Module 11 — Spending Analytics" in para.text:
            para.text = "Module 4 — Note System: A persistence layer allowing users to bookmark AI responses and manage personal study snippets contextually linked to subjects."
        elif "Module 5 — Wellness AI" in para.text:
            para.text = "Module 5 — Practice Hub: Generates AI-driven practice questions and tracks student scores to identify strengths and weaknesses."
        elif "Module 4 — Ordering" in para.text:
            para.text = "Module 6 — Revision Engine: Automates the creation of high-yield revision sheets from complex topics, focusing on formulas and exam-critical concepts."
        elif "Module 10 — Food Recommendation" in para.text:
             para.text = "" # Emptying out others
        elif "Module 2 — Gym Freak" in para.text:
             para.text = ""
        elif "Module 6 — Order Tracking" in para.text:
             para.text = ""
        elif "Module 8 — Feedback & Ratings" in para.text:
             para.text = ""
        elif "Module 9 — Global Admin" in para.text:
             para.text = ""

    # 2. Edit Tables
    if len(doc.tables) > 0:
        # Table 0: Student Contributions
        t0 = doc.tables[0]
        # Row 1 is the first student record
        t0.cell(1, 0).text = "CB.SC.U4CSE23144"
        t0.cell(1, 1).text = "Sandheep G S"
        t0.cell(1, 2).text = "Auth, Chat, Learning, Notes, Practice, Revision, Analytics"
        t0.cell(1, 3).text = (
            "Modules: 6 | Master Forms: 4 | Transaction Forms: 3 | "
            "Tables: 7 | Visualizations: 2 | Reports: 3 | Testing: E2E Auth & AI flow | API: Express REST API"
        )
        t0.cell(1, 4).text = "Khan Academy / GATE Overflow"
        # Clear other rows if they exist
        for i in range(2, len(t0.rows)):
            for j in range(len(t0.columns)):
                t0.cell(i, j).text = ""

    # Table 1: User Roles
    if len(doc.tables) > 1:
        t1 = doc.tables[1]
        t1.cell(1, 1).text = "Asking technical questions, saving notes, generating revision sheets, and tracking GATE preparation progress."
        # Removing Crew/Vendor/Admin for this specific simplified model or updating them
        t1.cell(2, 0).text = "Admin"
        t1.cell(2, 1).text = "Managing syllabus content and overseeing platform analytics."
        for i in range(3, len(t1.rows)):
            t1.cell(i, 0).text = ""
            t1.cell(i, 1).text = ""

    # Table 2: Comparison
    if len(doc.tables) > 2:
        t2 = doc.tables[2]
        comparisons = [
            ("Khan Academy", "Structured learning, quiz mode", "Inspired the structured subject explorer and practice hub."),
            ("GATE Overflow", "Question database, community solutions", "Inspired the technical focus and need for accurate LaTeX math rendering."),
            ("MyFitnessPal", "Goal tracking", "Inspired the progress analytics and mastery charts."),
            ("ChatGPT", "AI chat engine", "The core driver for the interactive tutor functionality.")
        ]
        for i, (prod, feat, infl) in enumerate(comparisons):
            if i + 1 < len(t2.rows):
                t2.cell(i+1, 1).text = prod
                t2.cell(i+1, 2).text = feat
                t2.cell(i+1, 3).text = infl

    # Table 3/4: Forms
    if len(doc.tables) > 3:
        t3 = doc.tables[3]
        forms = [("Signup Form", "Sandheep", "Auth"), ("Login Form", "Sandheep", "Auth"), ("Create Note", "Sandheep", "Notes"), ("Goal Set", "Sandheep", "Practice")]
        for i, (f, o, m) in enumerate(forms):
            if i + 1 < len(t3.rows):
                t3.cell(i+1, 1).text = f
                t3.cell(i+1, 2).text = o
                t3.cell(i+1, 3).text = m
    
    if len(doc.tables) > 4:
        t4 = doc.tables[4]
        tforms = [("Chat Input", "Sandheep", "Chat"), ("Search Notes", "Sandheep", "Notes"), ("Submit Quiz", "Sandheep", "Practice")]
        for i, (f, o, m) in enumerate(tforms):
            if i + 1 < len(t4.rows):
                t4.cell(i+1, 1).text = f
                t4.cell(i+1, 2).text = o
                t4.cell(i+1, 3).text = m

    doc.save(output_path)
    print(f"File saved successfully to {output_path}")

if __name__ == "__main__":
    edit_docx_r1("Review-1-Template (1).docx", "CB.SC.U4CSE23144Review 1_GateGPT.docx")
