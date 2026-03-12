from docx import Document
import sys

def analyze_docx(file_path):
    try:
        doc = Document(file_path)
        with open("docx_structure_r1.txt", "w", encoding="utf-8") as f:
            f.write("--- PARAGRAPHS ---\n")
            for i, para in enumerate(doc.paragraphs):
                f.write(f"P{i}: {para.text}\n")
            
            f.write("\n--- TABLES ---\n")
            for i, table in enumerate(doc.tables):
                f.write(f"Table {i}:\n")
                for r_idx, row in enumerate(table.rows):
                    row_text = [cell.text.strip().replace('\n', ' ') for cell in row.cells]
                    f.write(f"  Row {r_idx}: {' | '.join(row_text)}\n")
        print("Analysis written to docx_structure_r1.txt")
    except Exception as e:
        print(f"Error: {e}")

if __name__ == "__main__":
    analyze_docx("Review-1-Template (1).docx")
